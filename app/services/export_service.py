"""Export service — PDF and DOCX generation from resume data."""
import hashlib
import sys
from io import BytesIO

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from flask import render_template

# ── Python 3.8 / macOS OpenSSL compatibility patch ───────────────────────────
# reportlab 4.x calls hashlib.md5(usedforsecurity=False) which is only valid
# on Python 3.9+ (or standard CPython hashlib). On Python 3.8 with the macOS
# system OpenSSL, the keyword is rejected. We patch it away here before any
# reportlab import so the flag is silently ignored.
if sys.version_info < (3, 9):
    _orig_md5 = hashlib.md5

    def _patched_md5(*args, **kwargs):
        kwargs.pop("usedforsecurity", None)
        return _orig_md5(*args, **kwargs)

    hashlib.md5 = _patched_md5  # type: ignore[assignment]

PDF_ENGINE = None
_WeasyHTML = None

try:
    from xhtml2pdf import pisa
    PDF_ENGINE = "xhtml2pdf"
except ImportError:
    pass

try:
    from weasyprint import HTML as _WeasyHTML  # noqa: F401
    if PDF_ENGINE is None:
        PDF_ENGINE = "weasyprint"
except Exception:
    # weasyprint may be installed but missing system libraries (pango, cairo)
    pass


def build_pdf(data: dict) -> BytesIO:
    """Render resume HTML and convert to PDF bytes.

    Raises:
        RuntimeError: when no PDF engine is installed.
    """
    if PDF_ENGINE is None:
        raise RuntimeError(
            "No PDF engine installed. Run: pip install xhtml2pdf  OR  pip install weasyprint"
        )

    template_name = data.get("template", "classic")
    html_content = render_template(f"resume/{template_name}.html", **data)

    buffer = BytesIO()
    if PDF_ENGINE == "weasyprint":
        _WeasyHTML(string=html_content).write_pdf(buffer)
    else:
        pisa.CreatePDF(BytesIO(html_content.encode("utf-8")), dest=buffer, encoding="utf-8")
    buffer.seek(0)
    return buffer


def build_docx(data: dict) -> BytesIO:
    """Build a DOCX document from resume data and return as bytes."""
    contacts = data.get("contacts", {})
    name = contacts.get("name", "")
    job_title = contacts.get("jobTitle", "")

    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.27)
    section.page_height = Inches(11.69)

    # Header — name and job title
    name_paragraph = doc.add_paragraph()
    name_run = name_paragraph.add_run(name)
    name_run.bold = True
    name_run.font.size = Pt(18)
    name_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    if job_title:
        title_paragraph = doc.add_paragraph(job_title)
        title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    # Contact information
    contact_parts = [
        contacts[field]
        for field in ("email", "phone", "location")
        if contacts.get(field)
    ]
    if contact_parts:
        doc.add_paragraph(" | ".join(contact_parts))

    # Summary
    if data.get("summary"):
        doc.add_paragraph()
        doc.add_paragraph("SUMMARY", style="Heading 2")
        doc.add_paragraph(data["summary"])

    # Experience
    doc.add_paragraph()
    doc.add_paragraph("EXPERIENCE", style="Heading 2")
    for entry in data.get("experience", []):
        p = doc.add_paragraph()
        p.add_run(f"{entry.get('role', '')} — {entry.get('company', '')}").bold = True
        doc.add_paragraph(entry.get("dates", ""), style="List Bullet")
        doc.add_paragraph(entry.get("description", ""))

    # Education
    doc.add_paragraph()
    doc.add_paragraph("EDUCATION", style="Heading 2")
    for entry in data.get("education", []):
        p = doc.add_paragraph()
        p.add_run(entry.get("school", "")).bold = True
        doc.add_paragraph(f"{entry.get('degree', '')} · {entry.get('dates', '')}")

    # Skills
    if data.get("skills"):
        doc.add_paragraph()
        doc.add_paragraph("SKILLS", style="Heading 2")
        doc.add_paragraph(", ".join(data["skills"]))

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer
